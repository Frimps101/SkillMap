"""Plain-text extraction from uploaded CV files (PDF, DOCX, TXT)."""

import io
import re

ALLOWED_EXTENSIONS = (".pdf", ".docx", ".txt")
MAX_FILE_BYTES = 5 * 1024 * 1024  # 5 MB

# Control characters PostgreSQL/text processing can't handle (keep \n and \t)
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean(text: str) -> str:
    """Strip NUL and other control characters that break Postgres text columns."""
    text = _CONTROL_CHARS.sub(" ", text)
    return re.sub(r"[ \t]{2,}", " ", text).strip()


def extract_text(data: bytes, filename: str) -> str:
    name = filename.lower()

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        raw = "\n".join((page.extract_text() or "") for page in reader.pages)
        return _clean(raw)

    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return _clean("\n".join(parts))

    if name.endswith(".txt"):
        return _clean(data.decode("utf-8", errors="ignore"))

    raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT file.")
